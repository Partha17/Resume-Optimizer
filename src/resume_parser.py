"""Parse a resume PDF or DOCX into the structured dict consumed by the rest of the pipeline.

Two-stage strategy:

1. Raw text extraction via `pdfplumber` (PDF) or `python-docx` (DOCX) plus a fast
   heuristic section splitter — works fully offline and is good enough for
   well-formatted resumes.
2. If Gemini is configured, a second pass cleans up messy resumes (multi-column
   PDFs, mixed casing, oddly ordered sections) using `prompts/extract_resume.md`.

Both stages produce the same schema, so downstream code is oblivious to the path.
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from . import llm

EMPTY_RESUME: dict[str, Any] = {
    "contact": {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "linkedin": "",
        "github": "",
        "portfolio": "",
    },
    "summary": "",
    "experience": [],
    "skills": [],
    "education": [],
    "certifications": [],
    "projects": [],
    "publications": [],
    "languages": [],
    "awards": [],
}


# ---------- raw text extraction ----------

def _read_pdf(path: Path) -> str:
    import pdfplumber

    out: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(x_tolerance=1.5) or ""
            out.append(txt)
    return "\n".join(out)


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # also grab table cell text — some resumes use tables for layout
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def read_resume_text(path: str | Path) -> str:
    """Public helper: extract plain text from PDF or DOCX."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Resume not found: {p}")
    ext = p.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(p)
    if ext in {".docx", ".doc"}:
        return _read_docx(p)
    if ext == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume extension {ext!r}; use PDF, DOCX, or TXT")


# ---------- heuristic section split (offline fallback) ----------

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,5}\)?[\s.-]?)?\d{3,5}[\s.-]?\d{3,5}")
_LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+", re.I)
_GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w-]+", re.I)
_URL_RE = re.compile(r"https?://[^\s]+", re.I)

_SECTION_HEADERS = {
    "summary": ["summary", "professional summary", "profile", "objective", "about"],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "career history",
    ],
    "skills": ["skills", "technical skills", "core competencies", "key skills"],
    "education": ["education", "academic background", "qualifications"],
    "certifications": ["certifications", "certificates", "licenses"],
    "projects": ["projects", "key projects", "select projects"],
    "publications": ["publications", "papers", "research"],
    "languages": ["languages"],
    "awards": ["awards", "honors", "achievements", "accomplishments"],
}


def _detect_section(line: str) -> str | None:
    cleaned = line.strip().lower().rstrip(":")
    if len(cleaned) > 40:
        return None
    for canon, aliases in _SECTION_HEADERS.items():
        if cleaned in aliases:
            return canon
    return None


def _heuristic_parse(text: str) -> dict[str, Any]:
    """Best-effort offline parse. Always returns the full schema, possibly empty."""
    result: dict[str, Any] = json.loads(json.dumps(EMPTY_RESUME))  # deep copy
    lines = [ln.rstrip() for ln in text.splitlines()]

    # contact info — scrape the top 20 lines
    head = "\n".join(lines[:20])
    if m := _EMAIL_RE.search(head):
        result["contact"]["email"] = m.group(0)
    if m := _PHONE_RE.search(head):
        phone = re.sub(r"[^\d+]", "", m.group(0))
        if 7 <= len(phone.lstrip("+")) <= 15:
            result["contact"]["phone"] = m.group(0).strip()
    if m := _LINKEDIN_RE.search(text):
        result["contact"]["linkedin"] = m.group(0)
    if m := _GITHUB_RE.search(text):
        result["contact"]["github"] = m.group(0)

    # name guess: first non-empty line that isn't an email/url/phone
    for ln in lines[:10]:
        s = ln.strip()
        if not s:
            continue
        if _EMAIL_RE.search(s) or _URL_RE.search(s) or _PHONE_RE.search(s):
            continue
        # avoid lines with too many lowercase words (likely paragraphs)
        words = s.split()
        if 1 < len(words) <= 6 and sum(w[0].isupper() for w in words if w) >= len(words) - 1:
            result["contact"]["name"] = s
            break

    # group lines by section
    current = None
    buckets: dict[str, list[str]] = {k: [] for k in _SECTION_HEADERS}
    for ln in lines:
        sec = _detect_section(ln)
        if sec:
            current = sec
            continue
        if current and ln.strip():
            buckets[current].append(ln)

    if buckets["summary"]:
        result["summary"] = " ".join(b.strip() for b in buckets["summary"]).strip()

    if buckets["skills"]:
        raw = " ".join(buckets["skills"])
        # split on commas, semicolons, pipes, bullets
        parts = re.split(r"[,;|•·]| {2,}", raw)
        result["skills"] = sorted({p.strip(" -·•") for p in parts if 1 < len(p.strip()) < 60})

    if buckets["experience"]:
        result["experience"] = _parse_experience(buckets["experience"])

    if buckets["education"]:
        result["education"] = _parse_education(buckets["education"])

    if buckets["certifications"]:
        result["certifications"] = [
            {"name": ln.strip(" -•·"), "issuer": "", "date": ""}
            for ln in buckets["certifications"]
            if ln.strip()
        ]

    if buckets["projects"]:
        result["projects"] = _parse_projects(buckets["projects"])

    result["publications"] = [ln.strip(" -•·") for ln in buckets["publications"] if ln.strip()]
    result["languages"] = [ln.strip(" -•·") for ln in buckets["languages"] if ln.strip()]
    result["awards"] = [ln.strip(" -•·") for ln in buckets["awards"] if ln.strip()]

    return result


_DATE_RE = re.compile(
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{1,2}/\d{4}|\d{4})",
    re.I,
)
_DATE_RANGE_RE = re.compile(
    r"(?P<start>" + _DATE_RE.pattern + r")\s*[-–to]+\s*(?P<end>" + _DATE_RE.pattern + r"|Present|Current)",
    re.I,
)


def _parse_experience(lines: list[str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        is_bullet = s.startswith(("-", "*", "•", "·", "◦"))
        if is_bullet and current is not None:
            current["bullets"].append(s.lstrip(" -*•·◦").strip())
            continue

        m = _DATE_RANGE_RE.search(s)
        if m:
            # new role line: "Title — Company    Jan 2022 – Present"
            if current:
                entries.append(current)
            before = s[: m.start()].strip(" -–|,")
            title, company = (before.split(" at ", 1) + [""])[:2]
            if " — " in before or " - " in before:
                parts = re.split(r" [—–-] ", before, maxsplit=1)
                title, company = (parts + [""])[:2]
            elif "," in before:
                title, company = (before.split(",", 1) + [""])[:2]
            current = {
                "title": title.strip(),
                "company": company.strip(),
                "location": "",
                "start_date": m.group("start"),
                "end_date": m.group("end"),
                "bullets": [],
            }
        elif current is None:
            # first non-dated line — treat as a header without dates
            current = {
                "title": s,
                "company": "",
                "location": "",
                "start_date": "",
                "end_date": "",
                "bullets": [],
            }
        else:
            # continuation line; if no bullet marker, treat as a bullet
            current["bullets"].append(s)
    if current:
        entries.append(current)
    return entries


def _parse_education(lines: list[str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for ln in lines:
        s = ln.strip(" -•·")
        if not s:
            continue
        m = _DATE_RANGE_RE.search(s)
        rec = {
            "degree": "",
            "institution": "",
            "location": "",
            "start_date": m.group("start") if m else "",
            "end_date": m.group("end") if m else "",
            "details": "",
        }
        head = s[: m.start()].strip(" ,–-") if m else s
        if "," in head:
            rec["degree"], rec["institution"] = (head.split(",", 1) + [""])[:2]
            rec["degree"] = rec["degree"].strip()
            rec["institution"] = rec["institution"].strip()
        else:
            rec["degree"] = head
        entries.append(rec)
    return entries


def _parse_projects(lines: list[str]) -> list[dict[str, Any]]:
    projects: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        is_bullet = s.startswith(("-", "*", "•", "·"))
        if is_bullet and current:
            current["bullets"].append(s.lstrip(" -*•·").strip())
        else:
            if current:
                projects.append(current)
            current = {"name": s, "description": "", "bullets": []}
    if current:
        projects.append(current)
    return projects


# ---------- public API ----------

def parse_resume(path: str | Path, *, use_llm: bool = True) -> dict[str, Any]:
    """Parse a resume file into the canonical dict.

    Falls back to the heuristic parser if Gemini is not configured or fails.
    """
    text = read_resume_text(path)
    base = _heuristic_parse(text)

    if not use_llm or not llm.is_available():
        return base

    try:
        prompt = llm.load_prompt("extract_resume", resume_text=text[:30_000])
        cleaned = llm.complete_json(prompt)
    except Exception:
        return base

    return _merge(base, cleaned)


def _merge(base: dict[str, Any], llm_out: dict[str, Any]) -> dict[str, Any]:
    """Prefer LLM output where non-empty, fall back to heuristic. Preserves schema."""
    merged: dict[str, Any] = json.loads(json.dumps(EMPTY_RESUME))
    for key, default_val in EMPTY_RESUME.items():
        llm_val = llm_out.get(key, default_val)
        base_val = base.get(key, default_val)
        if isinstance(default_val, dict):
            merged[key] = {}
            llm_dict = llm_val if isinstance(llm_val, dict) else {}
            base_dict = base_val if isinstance(base_val, dict) else {}
            for sub, sub_default in default_val.items():
                merged[key][sub] = llm_dict.get(sub) or base_dict.get(sub) or sub_default
        elif isinstance(default_val, list):
            chosen = llm_val if (isinstance(llm_val, list) and llm_val) else base_val
            merged[key] = chosen if isinstance(chosen, list) else []
        else:
            merged[key] = llm_val or base_val or default_val
    return merged
