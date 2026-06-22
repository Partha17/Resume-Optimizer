"""Validate the optimised resume against ATS heuristics + measure keyword coverage.

Two score families:

1. **Keyword coverage** — fraction of critical / recommended keywords that appear
   in the rendered resume text. This is the lever that gets you to the top of
   recruiter search results.
2. **Format compliance** — parseable-headings, no images, no tables, no
   headers/footers, sane font count, reasonable page length.

`compare_before_after` returns a side-by-side DataFrame for inline notebook
display.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

from .keyword_extractor import KeywordRecord
from .resume_writer import SECTION_HEADINGS


_STANDARD_HEADINGS = {h.lower() for h in SECTION_HEADINGS.values()}


@dataclass
class ATSReport:
    keyword_coverage: dict[str, float] = field(default_factory=dict)
    keyword_hits: dict[str, list[str]] = field(default_factory=dict)
    keyword_misses: dict[str, list[str]] = field(default_factory=dict)
    format_checks: dict[str, bool] = field(default_factory=dict)
    format_notes: list[str] = field(default_factory=list)
    overall: float = 0.0

    def to_dict(self) -> dict[str, Any]:
        return {
            "keyword_coverage": self.keyword_coverage,
            "keyword_hits": self.keyword_hits,
            "keyword_misses": self.keyword_misses,
            "format_checks": self.format_checks,
            "format_notes": self.format_notes,
            "overall": self.overall,
        }


# ---------- keyword coverage ----------

def _normalise_text(t: str) -> str:
    return re.sub(r"\s+", " ", (t or "").lower())


def _contains(haystack: str, needle: str) -> bool:
    if not needle:
        return False
    # word-boundary check first; falls back to substring for hyphenated/acronym terms
    pattern = r"\b" + re.escape(needle.lower()) + r"\b"
    if re.search(pattern, haystack):
        return True
    return needle.lower() in haystack


def keyword_coverage(
    resume_text: str,
    records: list[KeywordRecord],
) -> tuple[dict[str, float], dict[str, list[str]], dict[str, list[str]]]:
    hay = _normalise_text(resume_text)
    coverage: dict[str, float] = {}
    hits: dict[str, list[str]] = {}
    misses: dict[str, list[str]] = {}

    for tier in ("critical", "recommended", "nice_to_have", "all"):
        if tier == "all":
            subset = records
        else:
            subset = [r for r in records if r.tier == tier]
        if not subset:
            coverage[tier] = 1.0  # vacuously satisfied
            hits[tier] = []
            misses[tier] = []
            continue
        present, absent = [], []
        for r in subset:
            if _contains(hay, r.keyword) or _contains(hay, r.display):
                present.append(r.display)
            else:
                absent.append(r.display)
        coverage[tier] = round(len(present) / len(subset), 4)
        hits[tier] = present
        misses[tier] = absent
    return coverage, hits, misses


# ---------- format compliance ----------

def _doc_to_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def format_checks(docx_path: Path) -> tuple[dict[str, bool], list[str]]:
    doc = Document(str(docx_path))
    notes: list[str] = []
    checks: dict[str, bool] = {}

    # tables
    checks["no_tables"] = len(doc.tables) == 0
    if not checks["no_tables"]:
        notes.append(f"Found {len(doc.tables)} table(s); some ATS parsers mis-read tabular layouts.")

    # images
    image_parts = [
        rel for rel in doc.part.rels.values() if "image" in rel.reltype  # type: ignore[attr-defined]
    ]
    checks["no_images"] = len(image_parts) == 0
    if not checks["no_images"]:
        notes.append(f"Found {len(image_parts)} image(s); strip them for ATS uploads.")

    # headers / footers
    has_header_footer = False
    for section in doc.sections:
        if section.header and any(p.text.strip() for p in section.header.paragraphs):
            has_header_footer = True
        if section.footer and any(p.text.strip() for p in section.footer.paragraphs):
            has_header_footer = True
    checks["no_headers_footers"] = not has_header_footer
    if has_header_footer:
        notes.append("Header/footer text found; many ATS strip these and lose contact info.")

    # headings
    text_lower = _doc_to_text(doc).lower()
    found_headings = [h for h in _STANDARD_HEADINGS if h in text_lower]
    checks["standard_headings"] = len(found_headings) >= 3
    if not checks["standard_headings"]:
        notes.append("Fewer than 3 standard section headings detected.")

    # font diversity
    fonts: set[str] = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fonts.add(run.font.name)
    checks["single_font"] = len(fonts) <= 1
    if len(fonts) > 1:
        notes.append(f"Multiple fonts in use: {sorted(fonts)}. Stick to one (Arial/Calibri).")

    # length (very rough — paragraph count proxy for pages)
    n_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    checks["reasonable_length"] = 15 <= n_paragraphs <= 120
    if n_paragraphs < 15:
        notes.append(f"Resume looks thin ({n_paragraphs} populated paragraphs); add detail.")
    elif n_paragraphs > 120:
        notes.append(
            f"Resume is long ({n_paragraphs} populated paragraphs); aim for ~1-2 pages."
        )

    # contact basics
    head_text = "\n".join(p.text for p in doc.paragraphs[:6])
    checks["has_email"] = bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", head_text))
    checks["has_phone"] = bool(re.search(r"\d[\d\s().+-]{6,}\d", head_text))
    if not checks["has_email"]:
        notes.append("Email not detected in the first 6 lines.")
    if not checks["has_phone"]:
        notes.append("Phone not detected in the first 6 lines.")

    return checks, notes


# ---------- top-level scorer ----------

def score_resume(
    docx_path: str | Path,
    market_records: list[KeywordRecord],
) -> ATSReport:
    docx_path = Path(docx_path)
    text = _doc_to_text(Document(str(docx_path)))

    coverage, hits, misses = keyword_coverage(text, market_records)
    fmt_checks, fmt_notes = format_checks(docx_path)

    # composite: 70% keywords (skewed to critical), 30% format
    # tier weights are renormalised across populated tiers so an empty tier
    # cannot vacuously inflate the score
    weights = {"critical": 0.6, "recommended": 0.3, "nice_to_have": 0.1}
    populated = {
        t: w for t, w in weights.items()
        if any(r.tier == t for r in market_records)
    }
    if populated:
        wsum = sum(populated.values())
        kw_score = sum(coverage.get(t, 0) * (w / wsum) for t, w in populated.items())
    else:
        kw_score = 0.0
    fmt_score = sum(1 for v in fmt_checks.values() if v) / max(len(fmt_checks), 1)
    overall = round(100 * (0.7 * kw_score + 0.3 * fmt_score), 1)

    return ATSReport(
        keyword_coverage={k: round(v * 100, 1) for k, v in coverage.items()},
        keyword_hits=hits,
        keyword_misses=misses,
        format_checks=fmt_checks,
        format_notes=fmt_notes,
        overall=overall,
    )


def score_against_jd(resume_text: str, jd_text: str, jd_keywords: list[str]) -> dict[str, Any]:
    """Lightweight per-JD coverage helper used when deciding to spin a tailored variant."""
    hay = _normalise_text(resume_text)
    if not jd_keywords:
        return {"coverage": 1.0, "missing": []}
    hits = [k for k in jd_keywords if _contains(hay, k)]
    missing = [k for k in jd_keywords if k not in hits]
    return {"coverage": round(len(hits) / len(jd_keywords), 4), "missing": missing}


def compare_before_after(
    before: ATSReport,
    after: ATSReport,
) -> pd.DataFrame:
    rows = [
        {
            "metric": f"Coverage: {tier}",
            "before %": before.keyword_coverage.get(tier, 0.0),
            "after %": after.keyword_coverage.get(tier, 0.0),
            "delta": round(
                after.keyword_coverage.get(tier, 0.0) - before.keyword_coverage.get(tier, 0.0),
                1,
            ),
        }
        for tier in ("critical", "recommended", "nice_to_have", "all")
    ]
    rows.append(
        {
            "metric": "Format checks passed",
            "before %": round(
                100 * sum(1 for v in before.format_checks.values() if v)
                / max(len(before.format_checks), 1),
                1,
            ),
            "after %": round(
                100 * sum(1 for v in after.format_checks.values() if v)
                / max(len(after.format_checks), 1),
                1,
            ),
            "delta": 0.0,
        }
    )
    rows[-1]["delta"] = rows[-1]["after %"] - rows[-1]["before %"]
    rows.append(
        {
            "metric": "Overall ATS score",
            "before %": before.overall,
            "after %": after.overall,
            "delta": round(after.overall - before.overall, 1),
        }
    )
    return pd.DataFrame(rows)
