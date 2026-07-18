"""Render a structured resume dict into an ATS-clean DOCX and orchestrate the
Gemini-driven rewrite that injects market keywords without fabrication.

ATS rules enforced by the writer:
  - Single column, no tables, no text boxes, no images, no headers/footers
  - Arial 11pt body, 12pt headings
  - Standard section headings: "Professional Summary", "Work Experience",
    "Skills", "Education", "Certifications", "Projects", "Publications"
  - Dates in "MMM YYYY" form
  - Hyphen bullet character (some ATS choke on Unicode bullets)
  - Skills as a comma-separated list (most ATS parsers prefer this to a bulleted list)
"""
from __future__ import annotations

import copy
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from . import llm
from .llm import LLMQuotaError

DEFAULT_FONT = "Arial"
BODY_SIZE = Pt(11)
HEADING_SIZE = Pt(12)
NAME_SIZE = Pt(18)
SECTION_ORDER = [
    "summary",
    "experience",
    "skills",
    "projects",
    "education",
    "certifications",
    "publications",
    "languages",
    "awards",
]
SECTION_HEADINGS = {
    "summary": "Professional Summary",
    "experience": "Work Experience",
    "skills": "Skills",
    "projects": "Projects",
    "education": "Education",
    "certifications": "Certifications",
    "publications": "Publications",
    "languages": "Languages",
    "awards": "Awards & Achievements",
}


# ---------- rewrite orchestration ----------

def optimize_resume(
    resume: dict[str, Any],
    *,
    role: str,
    domain: str,
    critical_keywords: list[str],
    recommended_keywords: list[str],
) -> dict[str, Any]:
    """Rewrite the summary and per-experience bullets using Gemini.

    If Gemini is not configured, returns the resume unchanged so the rest of the
    pipeline still runs end-to-end (you'll just see lower ATS coverage scores).
    """
    out = copy.deepcopy(resume)
    if not llm.is_available():
        print("[resume_writer] Gemini not configured; returning resume unchanged.")
        return out

    quota_hit = False

    if out.get("summary") or out.get("experience"):
        try:
            summary_prompt = llm.load_prompt(
                "rewrite_summary",
                role=role,
                domain=domain,
                current_summary=out.get("summary", ""),
                experience_json=json.dumps(out.get("experience", []), ensure_ascii=False)[:6000],
                skills_json=json.dumps(out.get("skills", []), ensure_ascii=False)[:2000],
                critical_keywords=json.dumps(critical_keywords[:25]),
                recommended_keywords=json.dumps(recommended_keywords[:25]),
            )
            new_summary = llm.complete(summary_prompt, temperature=0.4).strip()
            if new_summary:
                out["summary"] = new_summary
            llm.throttle(0.6)
        except LLMQuotaError as e:
            print(f"[resume_writer] quota hit on summary rewrite; aborting LLM stage: {str(e)[:160]}")
            quota_hit = True
        except Exception as e:
            print(f"[resume_writer] summary rewrite failed: {e}")

    if not quota_hit:
        for i, exp in enumerate(out.get("experience", [])):
            if not exp.get("bullets"):
                continue
            try:
                prompt = llm.load_prompt(
                    "rewrite_bullets",
                    role=role,
                    domain=domain,
                    experience_json=json.dumps(exp, ensure_ascii=False)[:4000],
                    critical_keywords=json.dumps(critical_keywords[:25]),
                    recommended_keywords=json.dumps(recommended_keywords[:25]),
                )
                new_bullets = llm.complete_json(prompt, temperature=0.4)
                if isinstance(new_bullets, list) and all(isinstance(b, str) for b in new_bullets):
                    out["experience"][i]["bullets"] = new_bullets
                llm.throttle(0.6)
            except LLMQuotaError as e:
                print(f"[resume_writer] quota hit on bullet rewrite (entry {i}); stopping: {str(e)[:160]}")
                quota_hit = True
                break
            except Exception as e:
                print(f"[resume_writer] bullet rewrite failed for entry {i}: {e}")
                continue

    # reorder skills so critical keywords come first
    out["skills"] = _reorder_skills(out.get("skills", []), critical_keywords, recommended_keywords)
    return out


def _reorder_skills(
    skills: list[str],
    critical: list[str],
    recommended: list[str],
) -> list[str]:
    if not skills:
        return skills
    crit_set = {k.lower() for k in critical}
    rec_set = {k.lower() for k in recommended}

    def bucket(s: str) -> int:
        sl = s.lower()
        if any(c in sl or sl in c for c in crit_set):
            return 0
        if any(c in sl or sl in c for c in rec_set):
            return 1
        return 2

    # preserve original order within each bucket
    indexed = list(enumerate(skills))
    indexed.sort(key=lambda pair: (bucket(pair[1]), pair[0]))
    return [s for _, s in indexed]


def tailor_for_job(
    master_resume: dict[str, Any],
    job: dict[str, Any],
    *,
    job_keywords: list[str],
) -> dict[str, Any]:
    """Per-job tailoring pass — adjusts summary/bullets/skills order only."""
    if not llm.is_available():
        return copy.deepcopy(master_resume)
    try:
        prompt = llm.load_prompt(
            "tailor_per_job",
            job_title=job.get("title", ""),
            company=job.get("company", ""),
            location=job.get("city", ""),
            job_description=(job.get("description") or "")[:6000],
            resume_json=json.dumps(master_resume, ensure_ascii=False)[:8000],
            job_keywords=json.dumps(job_keywords[:30]),
        )
        out = llm.complete_json(prompt, temperature=0.4)
        if isinstance(out, dict) and out.get("experience"):
            return out
    except Exception as e:
        print(f"[resume_writer] tailor failed for {job.get('company')}: {e}")
    return copy.deepcopy(master_resume)


# ---------- DOCX rendering ----------

_MONTHS = {
    "1": "Jan", "01": "Jan", "jan": "Jan", "january": "Jan",
    "2": "Feb", "02": "Feb", "feb": "Feb", "february": "Feb",
    "3": "Mar", "03": "Mar", "mar": "Mar", "march": "Mar",
    "4": "Apr", "04": "Apr", "apr": "Apr", "april": "Apr",
    "5": "May", "05": "May", "may": "May",
    "6": "Jun", "06": "Jun", "jun": "Jun", "june": "Jun",
    "7": "Jul", "07": "Jul", "jul": "Jul", "july": "Jul",
    "8": "Aug", "08": "Aug", "aug": "Aug", "august": "Aug",
    "9": "Sep", "09": "Sep", "sep": "Sep", "sept": "Sep", "september": "Sep",
    "10": "Oct", "oct": "Oct", "october": "Oct",
    "11": "Nov", "nov": "Nov", "november": "Nov",
    "12": "Dec", "dec": "Dec", "december": "Dec",
}


def _normalise_date(d: str) -> str:
    if not d:
        return ""
    s = d.strip()
    if s.lower() in {"present", "current", "now"}:
        return "Present"
    # already "MMM YYYY"?
    if re.match(r"^[A-Z][a-z]{2}\s\d{4}$", s):
        return s
    # "01/2022", "1/2022"
    if m := re.match(r"^(\d{1,2})[/-](\d{4})$", s):
        mon = _MONTHS.get(m.group(1).lstrip("0") or "1", "Jan")
        return f"{mon} {m.group(2)}"
    # "2022"
    if re.match(r"^\d{4}$", s):
        return s
    # "January 2022", "jan 2022", etc
    m = re.match(r"^(?P<mon>[A-Za-z]+)\.?\s+(?P<yr>\d{4})$", s)
    if m:
        mon = _MONTHS.get(m.group("mon").lower(), m.group("mon")[:3].title())
        return f"{mon} {m.group('yr')}"
    return s  # unrecognised; pass through


def _set_font(run, *, bold: bool = False, size: Pt = BODY_SIZE, color: tuple[int, int, int] = (0, 0, 0)) -> None:
    run.font.name = DEFAULT_FONT
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, bold=True, size=HEADING_SIZE)
    # underline via bottom border using the paragraph style is fiddly; bold uppercase is enough for ATS


def _contact_line(doc: Document, contact: dict[str, str]) -> None:
    name = contact.get("name") or ""
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(name)
        _set_font(run, bold=True, size=NAME_SIZE)

    contact_bits: list[str] = []
    for key in ("location", "phone", "email", "linkedin", "github", "portfolio"):
        v = contact.get(key)
        if v:
            contact_bits.append(v)
    if contact_bits:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(" | ".join(contact_bits))
        _set_font(run, size=Pt(10))


def _flatten(text: str) -> str:
    """Collapse PDF-extraction newlines and excess whitespace within a single bullet."""
    return re.sub(r"\s+", " ", text or "").strip()


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        flat = _flatten(item)
        if not flat:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"- {flat}")
        _set_font(run)


def _experience_block(doc: Document, exp: dict[str, Any]) -> None:
    title = exp.get("title") or ""
    company = exp.get("company") or ""
    location = exp.get("location") or ""
    start = _normalise_date(exp.get("start_date") or "")
    end = _normalise_date(exp.get("end_date") or "")
    date_range = " – ".join(d for d in [start, end] if d) if (start or end) else ""

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    # left: title at company / right: dates
    left_bits = [b for b in [title, company] if b]
    left = " | ".join(left_bits)
    run = p.add_run(left)
    _set_font(run, bold=True)
    if date_range:
        run = p.add_run(f"    {date_range}")
        _set_font(run, size=Pt(10))

    if location:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        run = p2.add_run(location)
        _set_font(run, size=Pt(10))

    _bullets(doc, exp.get("bullets") or [])


def _education_block(doc: Document, edu: dict[str, Any]) -> None:
    degree = edu.get("degree") or ""
    inst = edu.get("institution") or ""
    location = edu.get("location") or ""
    start = _normalise_date(edu.get("start_date") or "")
    end = _normalise_date(edu.get("end_date") or "")
    range_ = " – ".join(d for d in [start, end] if d) if (start or end) else ""

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    left_bits = [b for b in [degree, inst] if b]
    run = p.add_run(" | ".join(left_bits))
    _set_font(run, bold=True)
    if range_:
        run = p.add_run(f"    {range_}")
        _set_font(run, size=Pt(10))

    extras = " | ".join(b for b in [location, edu.get("details") or ""] if b)
    if extras:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        run = p2.add_run(extras)
        _set_font(run, size=Pt(10))


def _cert_block(doc: Document, cert: dict[str, Any]) -> None:
    line = " | ".join(
        b for b in [
            cert.get("name") or "",
            cert.get("issuer") or "",
            _normalise_date(cert.get("date") or ""),
        ] if b
    )
    if line:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"- {line}")
        _set_font(run)


def _project_block(doc: Document, proj: dict[str, Any]) -> None:
    name = proj.get("name") or ""
    if name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(name)
        _set_font(run, bold=True)
    if proj.get("description"):
        p = doc.add_paragraph()
        run = p.add_run(proj["description"])
        _set_font(run)
    _bullets(doc, proj.get("bullets") or [])


def write_docx(resume: dict[str, Any], output_path: str | Path) -> Path:
    """Render the resume dict to an ATS-clean .docx file."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Tighten default styles
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(2)

    _contact_line(doc, resume.get("contact") or {})

    for section in SECTION_ORDER:
        value = resume.get(section)
        if not value:
            continue

        if section == "summary":
            _heading(doc, SECTION_HEADINGS[section])
            p = doc.add_paragraph()
            run = p.add_run(value)
            _set_font(run)

        elif section == "experience":
            _heading(doc, SECTION_HEADINGS[section])
            for exp in value:
                _experience_block(doc, exp)

        elif section == "skills":
            _heading(doc, SECTION_HEADINGS[section])
            p = doc.add_paragraph()
            run = p.add_run(", ".join(value))
            _set_font(run)

        elif section == "projects":
            _heading(doc, SECTION_HEADINGS[section])
            for proj in value:
                _project_block(doc, proj)

        elif section == "education":
            _heading(doc, SECTION_HEADINGS[section])
            for edu in value:
                _education_block(doc, edu)

        elif section == "certifications":
            _heading(doc, SECTION_HEADINGS[section])
            for cert in value:
                _cert_block(doc, cert)

        elif section in {"publications", "languages", "awards"}:
            _heading(doc, SECTION_HEADINGS[section])
            for item in value:
                if isinstance(item, str) and item.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(14)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(f"- {item.strip()}")
                    _set_font(run)

    doc.save(str(output_path))
    return output_path


def write_pdf_if_possible(docx_path: str | Path, pdf_path: str | Path) -> Path | None:
    """Best-effort PDF render via docx2pdf (needs Word on mac/win or LibreOffice).

    Silently returns None if the conversion fails; the DOCX is still the canonical
    ATS upload anyway, so PDF is purely for human-eye polish.
    """
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx2pdf import convert  # type: ignore

        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except Exception as e:
        print(f"[resume_writer] PDF conversion skipped ({e}); DOCX is the canonical upload anyway.")
        return None


def safe_filename(text: str) -> str:
    """Slug helper for tailored variants."""
    text = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_").lower()
    return text[:40] or datetime.utcnow().strftime("%Y%m%d_%H%M%S")
