"""Per-job cover letter generator.

Produces a single-page DOCX cover letter grounded in the candidate's parsed
resume and the JD's top critical keywords. The prompt enforces no-fabrication
and lab-floor credibility; this module just stages inputs, calls Gemini, and
renders the DOCX.
"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from . import llm

DEFAULT_FONT = "Arial"
BODY_SIZE = Pt(11)


def _candidate_facts(resume: dict[str, Any]) -> dict[str, Any]:
    """Project the parsed resume down to the facts the prompt is allowed to use."""
    contact = resume.get("contact", {}) or {}
    return {
        "name": contact.get("name") or "",
        "location": contact.get("location") or "",
        "email": contact.get("email") or "",
        "phone": contact.get("phone") or "",
        "summary": (resume.get("summary") or "")[:1200],
        "skills": (resume.get("skills") or [])[:60],
        "experience": [
            {
                "title": e.get("title", ""),
                "company": e.get("company", ""),
                "location": e.get("location", ""),
                "start_date": e.get("start_date", ""),
                "end_date": e.get("end_date", ""),
                "bullets": (e.get("bullets") or [])[:8],
            }
            for e in (resume.get("experience") or [])[:6]
        ],
        "certifications": [
            c.get("name", "") for c in (resume.get("certifications") or [])
        ],
        "education": [
            {"degree": e.get("degree", ""), "institution": e.get("institution", "")}
            for e in (resume.get("education") or [])
        ],
    }


def generate_cover_letter(
    resume: dict[str, Any],
    job: dict[str, Any],
    *,
    role: str,
    domain: str,
    critical_keywords: list[str],
) -> str | None:
    """Call Gemini and return the cover-letter text, or None if the model is unconfigured."""
    if not llm.is_available():
        print("[cover_letter] Gemini not configured; skipping cover letter generation.")
        return None

    facts = _candidate_facts(resume)
    try:
        prompt = llm.load_prompt(
            "cover_letter",
            role=role,
            domain=domain,
            job_title=job.get("job_title") or job.get("title") or "",
            company=job.get("employer_name") or job.get("company") or "",
            location=job.get("job_city") or job.get("city") or "",
            job_description=(job.get("job_description") or job.get("description") or "")[:6000],
            candidate_facts=json.dumps(facts, ensure_ascii=False)[:8000],
            critical_keywords=json.dumps(list(critical_keywords)[:20]),
        )
        text = llm.complete(prompt, temperature=0.4).strip()
    except Exception as e:
        print(f"[cover_letter] generation failed: {e}")
        return None
    llm.throttle(0.5)
    return text or None


def write_cover_letter_docx(
    text: str,
    resume: dict[str, Any],
    job: dict[str, Any],
    output_path: str | Path,
) -> Path:
    """Render the cover letter to a clean single-page DOCX."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(6)

    contact = resume.get("contact") or {}
    name = contact.get("name") or ""
    bits = [contact.get(k) or "" for k in ("location", "phone", "email", "linkedin")]
    bits = [b for b in bits if b]

    # Header: name then contact line
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    if bits:
        p = doc.add_paragraph()
        run = p.add_run(" | ".join(bits))
        run.font.size = Pt(10)

    # Date
    p = doc.add_paragraph()
    p.add_run(datetime.now().strftime("%B %d, %Y")).font.size = BODY_SIZE

    # Recipient block
    company = job.get("employer_name") or job.get("company") or ""
    job_title = job.get("job_title") or job.get("title") or ""
    location = job.get("job_city") or job.get("city") or ""
    if company:
        p = doc.add_paragraph()
        p.add_run(f"Hiring Manager — {company}").bold = True
    if job_title:
        doc.add_paragraph(f"Role: {job_title}")
    if location:
        doc.add_paragraph(f"Location: {location}")
    doc.add_paragraph("")  # spacer

    # Body
    for para in text.split("\n"):
        para = para.strip()
        if not para:
            continue
        doc.add_paragraph(para)

    doc.save(str(output_path))
    return output_path


def generate_and_save(
    resume: dict[str, Any],
    job: dict[str, Any],
    output_path: str | Path,
    *,
    role: str,
    domain: str,
    critical_keywords: list[str],
) -> tuple[Path | None, str | None]:
    """End-to-end helper: generate text + write DOCX. Returns (path, text)."""
    text = generate_cover_letter(
        resume,
        job,
        role=role,
        domain=domain,
        critical_keywords=critical_keywords,
    )
    if not text:
        return None, None
    path = write_cover_letter_docx(text, resume, job, output_path)
    return path, text
